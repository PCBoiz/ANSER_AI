"""
src/core/doc_extract.py — bóc chữ từ file tài liệu. Code thuần, không LLM.

VÌ SAO TÁCH RA
--------------
`knowledge.py` bản cũ có `_read_file()` đọc theo đường dẫn trên đĩa. Bản viết lại
bỏ nó đi vì kho tri thức nhận `str` chứ không nhận file — nhưng thế thì không ai
bóc chữ ra nữa, và đường nạp tài liệu chưa bao giờ tồn tại.

Ở đây nhận BYTES, không nhận đường dẫn: tài liệu nội bộ của khách (hợp đồng,
bảng giá, chính sách công nợ) đi qua HTTP rồi bóc chữ trong bộ nhớ, không nằm lại
trên đĩa máy Brain (P2).

CHỖ HỎNG IM LẶNG PHẢI CHẶN
--------------------------
PDF SCAN không có lớp chữ. `pypdf` đọc nó ra chuỗi rỗng và **không báo lỗi gì**.
Nạp trót lọt, hiện "đã nạp 0 đoạn", rồi ba tuần sau người dùng hỏi về hợp đồng đó
và hệ thống trả lời "không có trong tài liệu" — trong khi file rành rành nằm trong
danh sách. Nên ở đây: bóc ra quá ít chữ so với số trang là BÁO ĐỘNG, không phải
lặng lẽ trả rỗng.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from io import BytesIO

logger = logging.getLogger("projecta.doc_extract")

# Đuôi file nhận được. `.doc` (Word 97) CỐ Ý không có: python-docx không đọc
# được, và im lặng trả rỗng thì tệ hơn từ chối thẳng.
SUPPORTED = (".pdf", ".docx", ".txt", ".md")

# Dưới ngần này ký tự mỗi trang thì gần như chắc chắn là PDF scan.
_MIN_CHARS_PER_PAGE = 50


@dataclass
class ExtractResult:
    text: str
    pages: int = 0
    warnings: list[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return bool(self.text.strip())


def _ext(filename: str) -> str:
    name = (filename or "").lower()
    return name[name.rfind("."):] if "." in name else ""


def _from_pdf(data: bytes) -> ExtractResult:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(f"Cần pypdf để đọc .pdf: {exc}") from exc

    reader = PdfReader(BytesIO(data))
    pages = len(reader.pages)
    text = "\n".join(p.extract_text() or "" for p in reader.pages)
    res = ExtractResult(text=text, pages=pages)

    if pages and len(text.strip()) < _MIN_CHARS_PER_PAGE * pages:
        res.warnings.append(
            f"Bóc được rất ít chữ ({len(text.strip())} ký tự trên {pages} trang) — "
            "gần như chắc chắn đây là PDF SCAN, không có lớp chữ. "
            "Xin bản mềm gốc (.docx), hoặc chạy OCR trước khi nạp."
        )
    return res


def _from_docx(data: bytes) -> ExtractResult:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(f"Cần python-docx để đọc .docx: {exc}") from exc

    d = docx.Document(BytesIO(data))
    phan = [p.text for p in d.paragraphs]

    # BẢNG trong .docx nằm NGOÀI `paragraphs`. Bảng giá cước của khách gần như
    # luôn là một bảng Word — đọc mỗi paragraphs là mất đúng phần giá trị nhất,
    # mà vẫn "nạp thành công" nên không ai biết.
    for bang in d.tables:
        for hang in bang.rows:
            o = [c.text.strip() for c in hang.cells]
            if any(o):
                phan.append("\t".join(o))

    return ExtractResult(text="\n".join(phan))


def extract(data: bytes, filename: str) -> ExtractResult:
    """
    Bóc chữ từ nội dung file. Ném `ValueError` khi không đọc được.

    Thông báo lỗi phải nói được PHẢI LÀM GÌ — người nạp tài liệu là kế toán hay
    chủ doanh nghiệp, không phải lập trình viên.
    """
    if not data:
        raise ValueError("File rỗng.")

    ext = _ext(filename)
    if ext not in SUPPORTED:
        raise ValueError(
            f"Chưa đọc được đuôi {ext or '(không rõ)'}. "
            f"Nhận: {', '.join(SUPPORTED)}. "
            "File .doc đời cũ thì mở bằng Word rồi Save As sang .docx."
        )

    try:
        if ext == ".pdf":
            res = _from_pdf(data)
        elif ext == ".docx":
            res = _from_docx(data)
        else:
            res = ExtractResult(text=data.decode("utf-8", errors="replace"))
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(
            f"Không mở được file ({type(exc).__name__}). File có thể hỏng hoặc "
            "tải lên dở chừng — thử tải lại."
        ) from exc

    if not res.ok:
        # Rỗng là LỖI, không phải "tài liệu không có gì". Nạp một tài liệu rỗng
        # rồi để nó nằm trong danh sách là tạo ra niềm tin sai: người dùng tưởng
        # hệ thống đã đọc hợp đồng đó.
        ly_do = res.warnings[0] if res.warnings else "File không chứa chữ nào đọc được."
        raise ValueError(ly_do)

    return res


__all__ = ["ExtractResult", "SUPPORTED", "extract"]
